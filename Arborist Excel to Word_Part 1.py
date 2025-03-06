#####################################################################################################
# Import Libraries, some libraries only imported later because importing now will intefere function #
#####################################################################################################
from openpyxl import load_workbook
from openpyxl import Workbook
import pandas as pd
from openpyxl.styles import PatternFill, Border, Side, Alignment, Protection, Font, Color, Fill
from openpyxl.workbook import Workbook
from openpyxl.reader.excel import load_workbook
import ospan
import sys
from docx.shared import RGBColor
from spire.doc.common import *
from docx.shared import Pt
from docx.shared import RGBColor
import docx 

#######################
# Determine variables #
#######################

#####################
# Working directory #
#####################
cwd = os.getcwd()

############
# filename #
############
filename = "Tree.xlsx"

################
# Photo folder #
################
try:
    narasimha = os.listdir(cwd + "/Photos")
except FileNotFoundError:
    print("Please rename your photo folder to Photos and ensure all photos are in their respective folders\nThank you")
    exit()

#########
# sheet #
#########
try:
    sheet = load_workbook(filename = filename)['Sheet1'] # Isolating the singular sheet containing the table
except KeyError:
    print("Sheet1 not found in " + filename + "\n"
        "Rename the sheet containing all the tables in " +
        str(filename) +
        " to Sheet1")
    exit()
except FileNotFoundError:
    print("The programme cannot find the Excel file containing the Tree table")
    print("Note that your working directory is:\n\n" + cwd + "\n")
    print("List of .xlsx file(s) in your working directory:" )
    excelfile = []
    for i in range(len(os.listdir(cwd))):
        if os.listdir(cwd)[i].endswith(".xlsx"):
            excelfile.append(os.listdir(cwd)[i])
    for i in range(len(excelfile)):
        print(str(i+1) + ") " + excelfile[i])
    print("\nPlease rename the .xlsx file to Tree.xlsx and try again")
    exit()

##########
# Others #
##########
row_count = sheet.max_row - 1
dest_filename = 'Xylopia.xlsx'
margin_size = 72.17


############################################################################################################################

#########################################################################
# From the Original List, split each row into its own sheet with header #
#########################################################################

print("Starting conversion from Excel to Word, this might take a while...\nPlease ensure that table columns in Tree.xlsx is arranged to Latitude, Longitude, Flora ID, Scientific Name, Common Name etc.")
print("Order of columns is important for code to work, columns after Common Name can be in any order you want")
print("\nIf a stream is dry, is it the streams fault?\nNo, it is the heavens fault for being dry") 
print("Sometimes when bad things happen, it is not our fault, but the work of Fate\n -Ramayana\n")
groupedRows = {}

#Group rows by 3rd column
for row_num, row in enumerate(sheet):
    rowData = [[cell.value, cell.style] for cell in row]
    if row_num != 0:
        if groupedRows.get(row[2].value, False):
            groupedRows[row[2].value].append(rowData)
        else:
            groupedRows[row[2].value] = [header]
            groupedRows[row[2].value].append(rowData)
    else:
        header = rowData

#Write each group to a sheet
wb = Workbook()
for group, rows in groupedRows.items():
    out_sheet = wb.create_sheet(title=group)
    for row_num, row in enumerate(rows):
        for col_num, col in enumerate(row):
            c = out_sheet.cell(row_num+1, col_num+1, col[0])
            c.style = col[1]

del wb['Sheet'] #Delete the first sheet since its blank for some reason

try:
    wb.save(filename = dest_filename)
except PermissionError:
    print("Do you have " + dest_filename + " open?\n" +
            "The code cannot edit an open file\n" +
            "Please close " + dest_filename + " before continuing\n")
    exit()
                                  
print(dest_filename + " has been printed, moving on to next phase\n")

############################################
# Truncate horizontal tables in each sheet #
############################################

print("Total Number of Trees Surveyed = " + str(row_count) + # Just to check how many trees we/you surveyed
        "\n" + 
        "Note: Total Trees Might Not Tally with Tree Number (T0XXX)")

try:
    for i in range(row_count): 
      df = pd.read_excel(dest_filename, i) # Isolating the sheet based on what value i is
      df1 = df.astype(object).T # Formula to Truncate
      Tree = "Tree" + str(i+1) + ".xlsx"
      df1.to_excel(Tree, 
        sheet_name = "Tree" + str(i+1), index=True, header=False) # Saving table into individual files Tree1, Tree2 etc 
      print("\n" + Tree + " has been created")
except PermissionError:
    print("\nDo you have " + Tree + " open?\n" + 
        "The code cannot edit an open file\n" +
        "Please close " + Tree + " before continuing\n")
    exit()

###########################################################
# Convert individual cell words to Aptos and Font size 12 #
###########################################################

for a in range(row_count):
    Tree = "Tree" + str(a+1) + ".xlsx"
    WB = load_workbook(Tree)
    varaha = WB["Tree"+str(a+1)]
    rc = varaha.max_row
    for i in range(rc):
        # Modifying all cells in first column
        _cells = varaha.cell(i+1,1)
        _cells.font = Font(name = "Aptos", size =12) 
        # Modifying all cells in second column
        _cells = varaha.cell(i+1,2) 
        _cells.font = Font(name = "Aptos", size = 12) 
    # Step 3: Italicise scientifc name by selecting cell containing scientific name
    _cell = varaha.cell(4, 2) 
    _cell.font = Font(name = "Aptos", i = True, size = 12)
    WB.save(Tree)
    print("\nFont for Values in " + Tree + " have been set to Aptos and Size 12")

##########################
# Left align cell values #
##########################

for i in range(row_count):
    Tree = "Tree" + str(i+1) + ".xlsx" 
    wb_style = load_workbook(Tree)
    sheet = wb_style.active
    Alstonia =sheet['A1':'B15']
    for cell in Alstonia:
        for c in cell:
            c.alignment=Alignment(horizontal='left', 
            vertical='center', shrinkToFit=False, textRotation=0, wrapText=True)
    wb_style.save(Tree)
    print("\nValues in " + Tree + " have been left alinged")

##########################     /  ##########################################
# Converts Excel to Word #    /   # Insert up to first 4 images into table # 
##########################   /    ##########################################

from spire.xls import *
from spire.doc import *


for i in range(row_count):
    try:
        #####################################
        # Establish variables for this loop #
        #####################################
        Tree = "Tree" + str(i+1) + ".xlsx" 
        Treedoc = "Tree" +str(i+1) + ".docx"
        workbook = Workbook()
        workbook.LoadFromFile(Tree)
        sheet = workbook.Worksheets[0]
        ROWS = sheet.Columns[0].RowCount
        doc = Document()

        ############################################
        # Creating the Word Document to imprint on #
        ############################################
        section = doc.AddSection()
        section.PageSetup.Orientation = PageOrientation.Portrait
        margins = section.PageSetup.Margins
        margins.Top = margin_size 
        margins.Bottom = margin_size
        margins.Left = margin_size
        margins.Right = margin_size
        para = section.AddParagraph()

        #######################################
        # Adding T0XXX Along with Common Name #
        #######################################
        df = pd.read_excel(Tree, sheet_name = "Tree" +str(i+1))
        # Find Tree number
        cell_value_a = df.iloc[1, 1]
        # Find Common name
        cell_value_b = df.iloc[3, 1]
        textRange = para.AppendText("           " +
                    str(cell_value_a) + 
                    " " + 
                    str(cell_value_b))

        ###########################################
        # Add and edit table to the Word Document #
        ###########################################
        table = section.AddTable(True)
          
        table.ResetCells(sheet.LastRow, sheet.LastColumn)
        table.ApplyHorizontalMerge(rowIndex= ROWS-1, startCellIndex= 0, endCellIndex= 1) # To merge the mysterious last row
                                                                                      # Can't find out where it come from
        for a in range(1,3): # Add a few more rows for merger
          table.AddRow() # Add one more row for photos
        table.ApplyVerticalMerge(columnIndex= 0, startRowIndex= ROWS-1, endRowIndex= ROWS+1) # Merge added rows vertically

        ####################################################################
        # Function to copy formatting from Excel cells to Word table cells #
        ####################################################################
        ###############################################################################################################
        # It replicates font properties, background color, and alignment (horizontal and vertical) from Excel to Word #
        ###############################################################################################################

        def copy_style(w_text_range, x_cell, w_cell):
             # Copy font formatting (color, size, font name, bold, italic) from Excel cell to Word text range
            w_text_range.CharacterFormat.TextColor = Color.FromRgb(x_cell.Style.Font.Color.R, x_cell.Style.Font.Color.G, x_cell.Style.Font.Color.B)
            w_text_range.CharacterFormat.FontSize = float(x_cell.Style.Font.Size)
            w_text_range.CharacterFormat.FontName = x_cell.Style.Font.FontName
            w_text_range.CharacterFormat.Bold = x_cell.Style.Font.IsBold
            w_text_range.CharacterFormat.Italic = x_cell.Style.Font.IsItalic

            # Copy background color from Excel cell to Word cell
            if x_cell.Style.FillPattern is not ExcelPatternType.none:
                w_cell.CellFormat.BackColor = Color.FromRgb(x_cell.Style.Color.R, x_cell.Style.Color.G, x_cell.Style.Color.B)

            # Copy horizontal alignment from Excel cell to Word paragraph
            if x_cell.HorizontalAlignment == HorizontalAlignType.Left:
                w_text_range.OwnerParagraph.Format.HorizontalAlignment = HorizontalAlignment.Left
            elif x_cell.HorizontalAlignment == HorizontalAlignType.Center:
                w_text_range.OwnerParagraph.Format.HorizontalAlignment = HorizontalAlignment.Center
            elif x_cell.HorizontalAlignment == HorizontalAlignType.Right:
                w_text_range.OwnerParagraph.Format.HorizontalAlignment = HorizontalAlignment.Right
              
            # Copy vertical alignment from Excel cell to Word cell
            if x_cell.VerticalAlignment == VerticalAlignType.Bottom:
                w_cell.CellFormat.VerticalAlignment = VerticalAlignment.Bottom
            elif x_cell.VerticalAlignment == VerticalAlignType.Center:
                w_cell.CellFormat.VerticalAlignment = VerticalAlignment.Middle
            elif x_cell.VerticalAlignment == VerticalAlignType.Top:
                w_cell.CellFormat.VerticalAlignment = VerticalAlignment.Top

        # Populate the Word table with data from the Excel sheet and copy the formatting
        for r in range(1, sheet.LastRow + 1):
            # Set row height in the Word table to match the row height in the Excel sheet
            table.Rows[r - 1].Height = float(sheet.Rows[r - 1].RowHeight)

            for c in range(1, sheet.LastColumn + 1):
                # Get the Excel cell at the current row and column
                x_cell = sheet.Range[r, c]

                 # Get the corresponding Word table cell
                w_cell = table.Rows[r - 1].Cells[c - 1]

                # Insert the text from the Excel cell into the Word table cell
                text_range = w_cell.AddParagraph().AppendText(x_cell.DisplayedText)

                # Copy the style (font, color, alignment, etc.) from the Excel cell to the Word cell
                copy_style(text_range, x_cell, w_cell)

        try:
            doc.SaveToFile(Treedoc, FileFormat.Docx2016) 
            print("\n" + Treedoc + " has been created")  
        except SpireException:
            print("\nDo you have " + Treedoc + " open?\n" + 
            "The code cannot edit an open file\n" +
            "Please close " + Treedoc + " before continuing\n")
            exit()

        #######################################     
        # Insert up to 4 photos into last row #
        #######################################
        workbook = Workbook()
        workbook.LoadFromFile(Tree)
        sheet = workbook.Worksheets[0]

        # Add an image to the 15 cell of the 1st row in the table
        dataf = pd.read_excel("Tree.xlsx", sheet_name = "Sheet1")
        cell_value_c = dataf.iloc[i, 2]
        cell = table.Rows[ROWS - 1].Cells[0]
        
        list = os.listdir(cwd + "/Photos" + 
                    "/" + 
                    str(cell_value_c))
        print("Tree" + str(i+1) + " ==" + str(cell_value_c))
        print("Photos in File" + " " + str(cell_value_c) + ": " + str(list))

        try:
          for r in range(4):
            picture = cell.Paragraphs[0].AppendPicture(
              cwd + "/Photos" +
              "/" + str(cell_value_c) + "/" + str(list[r])
                 )
            picture.Width = 150
            picture.Height = 150
            doc.SaveToFile(Treedoc, FileFormat.Docx2016)
            print("Photo " + str(list[r]) + " has been added to " + Treedoc + "(" + str(r+1) + " of 4" + ")")
        except IndexError:
          print("Only " + str(r) + " Photos Provided")
          pass

        # Save the generated Word document
        doc.SaveToFile(Treedoc, FileFormat.Docx2016) 
    except FileNotFoundError:
      print("No Photos Provided!")
      pass



#################################
# Replaces watermark with T0XXX #
#################################

from docx import Document

for i in range(row_count):
    #####################################
    # Establish variables for this loop #
    #####################################
    Tree = "Tree" +str(i+1) + ".xlsx"
    Treedoc = "Tree" + str(i+1) + ".docx"
    df = pd.read_excel(Tree, sheet_name = "Tree" +str(i+1))
    cell_value = df.iloc[1, 1]

    document = Document(Treedoc)
    style = document.styles["Normal"]
    font = style.font
    font.name = 'Aptos Display'
    font.size = Pt(20)
    font.color.rgb = RGBColor(15, 71, 97)
    replace_word = {'Evaluation Warning: The document was created with Spire.Doc for Python.': " "}
    for word in replace_word:
        for p in document.paragraphs:
            if p.text.find(word) >= 0:
                p.text = p.text.replace(word, replace_word[word])
    document.save(Treedoc)
    print("\nTitle Font, Colour and Size for " + Treedoc + " has been altered")

##################################################################################################
# Add page break to ensure each appended document starts on the next page (see next python file) #
##################################################################################################

for i in range(row_count):
    Treedoc = "Tree" + str(i+1) + ".docx" 
    doc = docx.Document(Treedoc)
    doc.add_page_break()
    doc.save(Treedoc)
    print("\nPage Break added to" + " " + Treedoc)